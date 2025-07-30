import PyPDF2
import pdfplumber
import aiofiles
import os
import docx
import docx2txt
from typing import Dict, Any
from app.config import settings
from app.services.ai_service import AIService

class DocumentService:
    def __init__(self):
        self.ai_service = AIService()
        self.upload_dir = os.path.join(settings.UPLOAD_DIR, "documents")

    async def save_uploaded_document(self, file_content: bytes, filename: str) -> str:
        """Save uploaded document file and return file path"""
        print(f"📁 [DOCUMENT SERVICE] Save document - filename: {filename}")
        print(f"📁 [DOCUMENT SERVICE] Upload directory: {self.upload_dir}")
        
        # Ensure upload directory exists
        if not os.path.exists(self.upload_dir):
            print(f"📁 [DOCUMENT SERVICE] Creating upload directory: {self.upload_dir}")
            os.makedirs(self.upload_dir, exist_ok=True)
        
        file_path = os.path.join(self.upload_dir, filename)
        print(f"📁 [DOCUMENT SERVICE] Full file path: {file_path}")
        
        try:
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            print(f"✅ [DOCUMENT SERVICE] File saved successfully: {file_path}")
        except Exception as e:
            print(f"❌ [DOCUMENT SERVICE] Failed to save file: {str(e)}")
            raise
        
        return file_path

    def extract_text_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"PyPDF2 extraction failed: {str(e)}")

    def extract_text_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber (more accurate)"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"pdfplumber extraction failed: {str(e)}")

    def extract_text_docx(self, file_path: str) -> str:
        """Extract text from Word documents using python-docx"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"python-docx extraction failed: {str(e)}")

    def extract_text_docx2txt(self, file_path: str) -> str:
        """Extract text from Word documents using docx2txt (fallback)"""
        try:
            text = docx2txt.process(file_path)
            return text.strip() if text else ""
        except Exception as e:
            raise Exception(f"docx2txt extraction failed: {str(e)}")

    def get_file_type(self, filename: str) -> str:
        """Determine file type from filename"""
        file_extension = filename.lower().split('.')[-1]
        if file_extension in ['pdf']:
            return 'pdf'
        elif file_extension in ['docx', 'doc']:
            return 'word'
        else:
            return 'unknown'

    async def process_resume_document(self, file_path: str, filename: str, user_id: str = None) -> Dict[str, Any]:
        """Process resume document (PDF, Word, etc.) and return structured data"""
        try:
            print(f"⚙️ [DOCUMENT SERVICE] Processing resume document: {file_path}")
            
            file_type = self.get_file_type(filename)
            print(f"📋 [DOCUMENT SERVICE] Detected file type: {file_type}")
            
            text_content = ""
            
            if file_type == 'pdf':
                # Try pdfplumber first (more accurate), fallback to PyPDF2
                try:
                    print("📖 [DOCUMENT SERVICE] Attempting PDF extraction with pdfplumber...")
                    text_content = self.extract_text_pdfplumber(file_path)
                    print("✅ [DOCUMENT SERVICE] pdfplumber extraction successful")
                except Exception as e:
                    print(f"⚠️ [DOCUMENT SERVICE] pdfplumber failed: {str(e)}")
                    print("📖 [DOCUMENT SERVICE] Falling back to PyPDF2...")
                    text_content = self.extract_text_pypdf2(file_path)
                    print("✅ [DOCUMENT SERVICE] PyPDF2 extraction successful")
                    
            elif file_type == 'word':
                # Try python-docx first, fallback to docx2txt
                try:
                    print("📖 [DOCUMENT SERVICE] Attempting Word extraction with python-docx...")
                    text_content = self.extract_text_docx(file_path)
                    print("✅ [DOCUMENT SERVICE] python-docx extraction successful")
                except Exception as e:
                    print(f"⚠️ [DOCUMENT SERVICE] python-docx failed: {str(e)}")
                    print("📖 [DOCUMENT SERVICE] Falling back to docx2txt...")
                    text_content = self.extract_text_docx2txt(file_path)
                    print("✅ [DOCUMENT SERVICE] docx2txt extraction successful")
            else:
                raise Exception(f"Unsupported file type: {file_type}")
            
            print(f"📄 [DOCUMENT SERVICE] Extracted text length: {len(text_content)} characters")
            
            # Process with AI regardless of content length - extract whatever is available
            print("🤖 [DOCUMENT SERVICE] Processing content with AI service...")
            
            # If no text content, return empty but successful result
            if not text_content or len(text_content.strip()) == 0:
                print("⚠️ [DOCUMENT SERVICE] No text content extracted, returning empty data")
                return {
                    "success": True,
                    "extracted_data": {},
                    "raw_text_length": 0
                }
            
            structured_data = await self.ai_service.process_resume_content(text_content, user_id)
            print(f"✅ [DOCUMENT SERVICE] AI processing completed")
            
            # Always return success, even if AI processing has issues
            if "error" in structured_data:
                print(f"⚠️ [DOCUMENT SERVICE] AI processing had issues: {structured_data['error']}")
                print("✅ [DOCUMENT SERVICE] Returning empty data but marking as successful")
                return {
                    "success": True,
                    "extracted_data": {},
                    "raw_text_length": len(text_content),
                    "processing_note": "AI processing encountered issues but document was processed"
                }
            
            print("✅ [DOCUMENT SERVICE] Document processing completed successfully")
            return {
                "success": True,
                "extracted_data": structured_data,
                "raw_text_length": len(text_content)
            }
            
        except Exception as e:
            print(f"❌ [DOCUMENT SERVICE] Exception during document processing: {str(e)}")
            print(f"❌ [DOCUMENT SERVICE] Exception type: {type(e).__name__}")
            return {
                "success": False,
                "error": f"Document processing failed: {str(e)}"
            }
        finally:
            # Clean up uploaded file
            if os.path.exists(file_path):
                os.remove(file_path)

    # Keep the old method for backward compatibility
    async def process_linkedin_pdf(self, file_path: str) -> Dict[str, Any]:
        """Legacy method - use process_resume_document instead"""
        return await self.process_resume_document(file_path, "document.pdf")