#!/usr/bin/env python3
"""
Test script to verify Word document extraction
"""
import asyncio
import sys
import os
from docx import Document

# Add the app directory to the Python path
sys.path.append(os.path.join(os.path.dirname(__file__), 'app'))

from app.services.pdf_service import DocumentService

def create_sample_word_doc():
    """Create a sample Word document for testing"""
    doc = Document()
    
    # Add title
    doc.add_heading('Jane Smith', 0)
    doc.add_paragraph('Marketing Manager | Digital Marketing Specialist')
    doc.add_paragraph('New York, NY | jane.smith@email.com | (555) 987-6543')
    doc.add_paragraph('LinkedIn: linkedin.com/in/janesmith')
    
    # Add summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('Results-driven marketing professional with 7+ years of experience in digital marketing, '
                     'content strategy, and brand management. Proven track record of increasing brand awareness '
                     'and driving customer engagement through innovative marketing campaigns.')
    
    # Add skills
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph('• Digital Marketing, SEO/SEM, Content Marketing')
    doc.add_paragraph('• Google Analytics, HubSpot, Salesforce')
    doc.add_paragraph('• Adobe Creative Suite, Canva, Figma')
    doc.add_paragraph('• Social Media Management, Email Marketing')
    
    # Add experience
    doc.add_heading('Work Experience', level=1)
    
    doc.add_heading('Senior Marketing Manager | GrowthCorp | 2021 - Present', level=2)
    doc.add_paragraph('• Led digital marketing initiatives resulting in 150% increase in online engagement')
    doc.add_paragraph('• Managed marketing budget of $500K+ and team of 5 marketing specialists')
    doc.add_paragraph('• Implemented data-driven marketing strategies increasing conversion rates by 35%')
    
    doc.add_heading('Marketing Specialist | BrandHub | 2018 - 2021', level=2)
    doc.add_paragraph('• Created and executed content marketing campaigns across multiple channels')
    doc.add_paragraph('• Analyzed market trends and competitor strategies to optimize marketing approach')
    doc.add_paragraph('• Collaborated with sales team to align marketing and sales objectives')
    
    # Add education
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Master of Business Administration (MBA)')
    doc.add_paragraph('New York University Stern School of Business | 2016 - 2018')
    doc.add_paragraph('')
    doc.add_paragraph('Bachelor of Arts in Marketing')
    doc.add_paragraph('Columbia University | 2012 - 2016')
    doc.add_paragraph('Magna Cum Laude, GPA: 3.7/4.0')
    
    # Add certifications
    doc.add_heading('Certifications', level=1)
    doc.add_paragraph('• Google Ads Certified Professional')
    doc.add_paragraph('• HubSpot Content Marketing Certification')
    doc.add_paragraph('• Facebook Blueprint Certification')
    
    # Save the document
    doc_path = os.path.join(os.path.dirname(__file__), 'sample_resume.docx')
    doc.save(doc_path)
    return doc_path

async def test_word_extraction():
    """Test Word document extraction"""
    
    print("🚀 Testing Word Document Extraction")
    print("=" * 50)
    
    # Create sample Word document
    print("📝 Creating sample Word document...")
    doc_path = create_sample_word_doc()
    print(f"✅ Sample document created: {doc_path}")
    
    # Test document service
    document_service = DocumentService()
    
    print("🔍 Testing Word document extraction...")
    result = await document_service.process_resume_document(doc_path, "sample_resume.docx")
    
    if not result.get('success'):
        print(f"❌ Document extraction failed: {result.get('error', 'Unknown error')}")
        return
    
    print("✅ Word document extraction successful!")
    
    extracted_data = result.get('extracted_data', {})
    
    print("\n📊 Extracted Data Summary:")
    print(f"Name: {extracted_data.get('name', 'Not found')}")
    print(f"Profession: {extracted_data.get('profession', 'Not detected')}")
    print(f"Designation: {extracted_data.get('designation', 'Not found')}")
    print(f"Experience: {extracted_data.get('experience', 'Not calculated')}")
    print(f"Skills count: {len(extracted_data.get('skills', []))}")
    print(f"Experience details count: {len(extracted_data.get('experience_details', []))}")
    print(f"Education count: {len(extracted_data.get('education', []))}")
    print(f"Certifications count: {len(extracted_data.get('certifications', []))}")
    print(f"Contact info available: {'Yes' if extracted_data.get('contact_info') else 'No'}")
    
    # Display some detailed examples
    if extracted_data.get('skills'):
        print(f"\n🔧 Sample Skills:")
        for skill in extracted_data.get('skills', [])[:3]:  # Show first 3
            print(f"  • {skill.get('name', 'Unknown')} - {skill.get('level', 'Unknown')} ({skill.get('years', 0)} years)")
    
    if extracted_data.get('experience_details'):
        print(f"\n💼 Sample Experience:")
        for exp in extracted_data.get('experience_details', [])[:2]:  # Show first 2
            print(f"  • {exp.get('position', 'Unknown')} at {exp.get('company', 'Unknown')} ({exp.get('duration', 'Unknown')})")
    
    if extracted_data.get('education'):
        print(f"\n🎓 Education:")
        for edu in extracted_data.get('education', []):
            print(f"  • {edu.get('degree', 'Unknown')} from {edu.get('institution', 'Unknown')}")
    
    # Clean up
    if os.path.exists(doc_path):
        os.remove(doc_path)
        print(f"\n🧹 Cleaned up test file: {doc_path}")
    
    print("\n✅ Word document extraction test completed successfully!")

if __name__ == "__main__":
    # Check if we have the required environment variables
    if not os.getenv('OPENAI_API_KEY'):
        print("❌ Error: OPENAI_API_KEY environment variable not set")
        print("Please set your OpenAI API key to test the extraction")
        sys.exit(1)
    
    # Run the test
    asyncio.run(test_word_extraction())